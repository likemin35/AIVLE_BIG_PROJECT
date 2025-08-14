import os
import json
import logging
import uuid
import urllib.parse
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS, cross_origin
import requests

# Vertex AI / LangChain / Chroma imports
try:
    import vertexai
    from vertexai.generative_models import GenerativeModel
    from google.oauth2 import service_account
    from google.cloud import secretmanager
    from langchain_community.document_loaders import PyPDFDirectoryLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_huggingface.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import Chroma
except Exception as e:
    # When running in environments without these packages, the app will still start
    # but certain endpoints will return an error explaining missing deps.
    vertexai = None
    GenerativeModel = None
    service_account = None
    secretmanager = None
    PyPDFDirectoryLoader = None
    RecursiveCharacterTextSplitter = None
    HuggingFaceEmbeddings = None
    Chroma = None
    missing_imports_exception = e

# DOCX utilities
try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

# ---------------------------
# App config
# ---------------------------
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# External services (can be configured via env)
TERM_SERVICE_URL = os.environ.get("TERM_SERVICE_URL", "http://localhost:8083/terms")
POINT_SERVICE_URL = os.environ.get("POINT_SERVICE_URL", "http://localhost:8085")

# GCP / Vertex config
PROJECT_ID = os.environ.get("PROJECT_ID", "aivle-team0721")
LOCATION = os.environ.get("LOCATION", "us-central1")
LOCAL_KEY_FILE = os.environ.get("LOCAL_KEY_FILE", os.path.join(os.path.dirname(__file__), "firebase-adminsdk.json"))

# Vector DB base directories (you can point these to your mounted drive or local path)
VECTOR_BASE_DIR = os.environ.get("VECTOR_BASE_DIR", os.path.join(os.path.dirname(__file__), "vector_dbs"))

# Expected category -> folder mapping (defaults can be overridden by env)
VECTOR_DB_MAP = {
    'loan': os.path.join(VECTOR_BASE_DIR, '대출'),
    'cancer_insurance': os.path.join(VECTOR_BASE_DIR, '암보험'),
    'deposit': os.path.join(VECTOR_BASE_DIR, '예금'),
    'car_insurance': os.path.join(VECTOR_BASE_DIR, '자동차보험'),
    'savings': os.path.join(VECTOR_BASE_DIR, '적금')
}

# 법령 벡터 DB (새로 추가)
LAW_VECTOR_DIR = os.environ.get("LAW_VECTOR_DIR", os.path.join(VECTOR_BASE_DIR, '법령'))

# Output folder for generated DOCX files
OUTPUT_DIR = os.environ.get("OUTPUT_DIR", os.path.join(os.path.dirname(__file__), 'outputs'))
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Prompt template (keeps original but trimmed for readability)
PROMPT_TEMPLATE = os.environ.get("PROMPT_TEMPLATE") or """
기업 이름은 다음과 같아:
{company_name}

상품 이름은 다음과 같아:
{product_name}

다음은 기업이 제공한 상품 정보야:
{wishlist}

다음은 해당 약관/계약의 시행 날짜야:
{date}

아래는 참고용 약관 문서 및 법령(참고자료)야:
{context}

위의 상품 정보와 약관/법령 문서를 참고해서 이 상품에 맞는 보험 약관 초안을 자세하게 작성해줘.
- 조항은 '제1조'... '1.' 형식으로 표기.
- 최소 50조항 이상 작성.
- 보장 관련 금액과 제외항목을 구체적으로 작성.
- 표(지급기준표나 해약환급금)를 JSON으로 따로 출력하고, 응답 최종에는 "TABLES_JSON_BEGIN"과 "TABLES_JSON_END" 사이에 JSON을 넣어줘.
- 추가로 DOCX로 저장하기에 적절한 구조(JSON)를 반환해줘.
"""

# ---------------------------
# Initialize Vertex AI, embeddings and model
# ---------------------------
credentials = None
gemini_model = None
embedding = None

if vertexai is None:
    logger.warning("Vertex AI or dependent libraries are missing. AI functions will be disabled.")
else:
    try:
        # Try Secret Manager first (same as original)
        secret_client = secretmanager.SecretManagerServiceClient()
        secret_name = f"projects/{PROJECT_ID}/secrets/firebase-adminsdk/versions/latest"
        response = secret_client.access_secret_version(name=secret_name)
        secret_payload = response.payload.data.decode("UTF-8")
        credentials_info = json.loads(secret_payload)
        credentials = service_account.Credentials.from_service_account_info(credentials_info)
        logger.info("Secret Manager에서 서비스 계정 키 로드 성공")
    except Exception as e:
        logger.warning(f"Secret Manager 접근 실패: {e}. 로컬 키 파일로 대체합니다.")
        try:
            if not os.path.exists(LOCAL_KEY_FILE):
                raise FileNotFoundError("로컬 서비스 계정 키 파일을 찾을 수 없습니다: " + LOCAL_KEY_FILE)
            credentials = service_account.Credentials.from_service_account_file(LOCAL_KEY_FILE)
            logger.info("로컬 파일에서 서비스 계정 키 로드 성공")
        except Exception as file_e:
            logger.error(f"AI 서비스 초기화 실패: Secret Manager와 로컬 파일 모두 실패. ({file_e})")
            credentials = None

if credentials and vertexai is not None:
    try:
        vertexai.init(project=PROJECT_ID, location=LOCATION, credentials=credentials)
        gemini_model = GenerativeModel("gemini-2.5-flash-lite")
        embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logger.info("Vertex AI 및 임베딩 초기화 성공")
    except Exception as e:
        logger.error(f"Vertex AI 또는 언어 모델 초기화 실패: {e}")
        gemini_model = None

# ---------------------------
# Helper functions
# ---------------------------

def call_point_service_reduce(user_id: str, amount: int, reason: str) -> (bool, dict):
    """Call point service to reduce points. Returns (ok, response_json_or_text)."""
    try:
        encoded_reason = urllib.parse.quote(reason)
        base_url = POINT_SERVICE_URL.rstrip('/')
        url = f"{base_url}/api/points/{user_id}/reduce?amount={amount}&reason={encoded_reason}"
        logger.info(f"포인트 차감 요청: {url}")
        r = requests.post(url)
        if r.ok:
            try:
                return True, r.json()
            except Exception:
                return True, {"raw": r.text}
        else:
            try:
                return False, r.json()
            except Exception:
                return False, {"raw": r.text}
    except requests.exceptions.RequestException as e:
        logger.exception("Point 서비스 호출 실패 (네트워크 오류)")
        return False, {"error": "Point 서비스 연결 실패"}


def load_context_from_vectorstores(category: str, wishlist: str, top_k: int = 5) -> str:
    """Load relevant docs from both product-specific vectorstore and 법령 vectorstore and combine them.
    Returns combined text context.
    """
    if Chroma is None or embedding is None:
        raise RuntimeError("Vectorstore or embedding library not initialized.")

    texts = []
    # product category
    db_dir = VECTOR_DB_MAP.get(category)
    if db_dir and os.path.isdir(db_dir):
        try:
            vs = Chroma(persist_directory=db_dir, embedding_function=embedding)
            retriever = vs.as_retriever(search_kwargs={"k": top_k})
            docs = retriever.invoke(wishlist)
            texts.extend([d.page_content for d in docs])
        except Exception:
            logger.exception(f"카테고리 벡터스토어 로드 실패: {db_dir}")

    # 법령 벡터
    if os.path.isdir(LAW_VECTOR_DIR):
        try:
            law_vs = Chroma(persist_directory=LAW_VECTOR_DIR, embedding_function=embedding)
            law_retriever = law_vs.as_retriever(search_kwargs={"k": top_k})
            law_docs = law_retriever.invoke(wishlist)
            texts.extend(["[법령] " + d.page_content for d in law_docs])
        except Exception:
            logger.exception("법령 벡터스토어 로드 실패")

    # fallback: if no texts, return empty string
    return "\n\n".join(texts)


def generate_terms_with_gemini(prompt: str) -> dict:
    """Call Gemini and return parsed result.
    Expect the model to place a JSON between TABLES_JSON_BEGIN and TABLES_JSON_END that contains structured tables.
    Return dict: {"text": ..., "tables": [...], "raw": full_response}
    """
    if gemini_model is None:
        raise RuntimeError("AI 모델이 초기화되지 않았습니다.")

    response = gemini_model.generate_content(prompt)
    # This structure follows previous usage; adapt if the SDK differs.
    try:
        generated_text = response.candidates[0].content.parts[0].text
    except Exception:
        generated_text = str(response)

    # Extract tables JSON if model followed the instruction
    tables = []
    TABLE_BEGIN = "TABLES_JSON_BEGIN"
    TABLE_END = "TABLES_JSON_END"
    if TABLE_BEGIN in generated_text and TABLE_END in generated_text:
        try:
            start = generated_text.index(TABLE_BEGIN) + len(TABLE_BEGIN)
            end = generated_text.index(TABLE_END)
            json_blob = generated_text[start:end].strip()
            tables = json.loads(json_blob)
            # remove the JSON segment from returned text (so `terms` is clean)
            cleaned_text = (generated_text[:generated_text.index(TABLE_BEGIN)] + generated_text[end + len(TABLE_END):]).strip()
        except Exception:
            logger.exception("모델이 반환한 TABLE JSON 파싱 실패")
            cleaned_text = generated_text
    else:
        cleaned_text = generated_text

    return {"text": cleaned_text, "tables": tables, "raw": generated_text}


def save_terms_to_docx(terms_text: str, tables: list, meta: dict) -> str:
    """Create a DOCX file including terms and tables. Returns filename (relative to OUTPUT_DIR).
    This function uses python-docx.
    """
    if Document is None:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다.")

    doc = Document()
    # Title
    title = meta.get("productName") or "약관 초안"
    p = doc.add_heading(title, level=1)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Meta info
    doc.add_paragraph(f"회사명: {meta.get('companyName', '')}")
    doc.add_paragraph(f"카테고리: {meta.get('category', '')}")
    doc.add_paragraph(f"작성일: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.add_paragraph("")

    # Terms body (large block)
    for line in terms_text.split('\n'):
        # add line as paragraph; preserve blank lines
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    # Insert tables (if any)
    if tables:
        doc.add_page_break()
        doc.add_heading("지급기준표 / 표", level=2)
        for idx, t in enumerate(tables):
            # Expecting table structure like: {"title":"...","headers":[...],"rows":[[...],[...]]}
            title = t.get('title', f'표 {idx+1}')
            headers = t.get('headers', [])
            rows = t.get('rows', [])
            doc.add_paragraph(title)
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                for r in rows:
                    row_cells = table.add_row().cells
                    for i, v in enumerate(r):
                        row_cells[i].text = str(v)
            else:
                # If no headers, just dump rows
                for r in rows:
                    doc.add_paragraph("\t".join([str(x) for x in r]))
            doc.add_paragraph("")

    # Save file
    filename = f"terms_{meta.get('companyName','unknown')}_{meta.get('productName','product')}_{uuid.uuid4().hex[:8]}.docx"
    # sanitize filename
    filename = filename.replace(' ', '_')
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    logger.info(f"DOCX 저장 완료: {filepath}")
    return filename


# ---------------------------
# API endpoints
# ---------------------------

@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,x-authenticated-user-uid')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response


@app.route('/api/generate', methods=['POST', 'OPTIONS'])
@cross_origin(origin='*')
def generate_terms():
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200

    # Basic dependency checks
    if vertexai is None or gemini_model is None:
        return jsonify({"error": "AI 라이브러리 또는 모델이 초기화되지 않았습니다. 서버 로그를 확인하세요."}), 500

    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "요청 데이터가 없습니다."}), 400

        company_name = data.get('companyName')
        category = data.get('category')
        product_name = data.get('productName')
        wishlist = data.get('requirements')
        user_id = request.headers.get('x-authenticated-user-uid')
        effective_date = data.get('effectiveDate')

        if not all([company_name, category, product_name, wishlist, user_id, effective_date]):
            logger.warning("필수 입력값 누락")
            return jsonify({"error": "필수 입력값이 누락되었습니다."}), 400

        # 1) 포인트 차감(기존 로직 유지)
        deduction_amount = int(os.environ.get('DEDUCTION_AMOUNT', 5000))
        ok, resp = call_point_service_reduce(user_id, deduction_amount, "AI 약관 초안 생성")
        if not ok:
            return jsonify({"error": resp}), 400

        # 2) 벡터 스토어에서 관련 문서 로드 (제품 + 법령)
        try:
            context = load_context_from_vectorstores(category, wishlist)
        except Exception as e:
            logger.exception("벡터스토어 로드 중 오류")
            return jsonify({"error": "벡터스토어 로드 중 오류가 발생했습니다."}), 500

        # 3) 프롬프트 조합 및 모델 호출
        prompt = PROMPT_TEMPLATE.format(
            context=context,
            company_name=company_name,
            product_name=product_name,
            wishlist=wishlist,
            date=effective_date
        )

        try:
            gen = generate_terms_with_gemini(prompt)
            generated_text = gen.get('text', '')
            tables = gen.get('tables', [])
        except Exception:
            logger.exception("모델 호출 중 오류")
            return jsonify({"error": "모델 호출 중 오류가 발생했습니다."}), 500

        # 4) DOCX 저장 (선택사항: 프론트에서 옵션으로 요청할 수 있음)
        docx_filename = None
        try:
            if data.get('saveDocx', True):
                meta = {"companyName": company_name, "category": category, "productName": product_name}
                docx_filename = save_terms_to_docx(generated_text, tables, meta)
        except Exception:
            logger.exception("DOCX 저장 중 오류")
            # DOCX 실패는 치명적이지 않으므로 계속 진행
            docx_filename = None

        result = {
            "terms": generated_text,
            "tables": tables,
            "meta": {
                "companyName": company_name,
                "category": category,
                "productName": product_name,
                "requirements": wishlist,
                "effectiveDate": effective_date
            }
        }

        if docx_filename:
            download_url = request.host_url.rstrip('/') + f"/api/download/{urllib.parse.quote(docx_filename)}"
            result['docxUrl'] = download_url

        return jsonify(result), 200

    except Exception:
        logger.exception("약관 생성 중 오류 발생")
        return jsonify({"error": "약관 생성 중 서버에서 오류가 발생했습니다."}), 500


@app.route('/api/download/<path:filename>', methods=['GET'])
def download_file(filename):
    # Serve files from OUTPUT_DIR
    safe_name = os.path.basename(filename)
    return send_from_directory(OUTPUT_DIR, safe_name, as_attachment=True)


# Health check
@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "ai_initialized": gemini_model is not None}), 200


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)), debug=True)