# c:\Users\User\Desktop\BIGProject\AIVLE_BIG_PROJECT\analyze_ai\analyze_terms.py
# analyze_terms.py
from __future__ import annotations

import os, re, json, uuid, logging, sys
from datetime import datetime, timezone
from typing import List, Tuple

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS, cross_origin
from werkzeug.utils import secure_filename

# --- Chroma 버전 로깅(디버그용) ---
try:
    import chromadb  # noqa
    logging.basicConfig(level=logging.INFO)
    logging.info(f"[BOOT] chromadb version: {chromadb.__version__}")
except Exception:
    pass

# Google / Vertex
import vertexai
from google.oauth2 import service_account
from google.cloud import secretmanager

# LangChain / RAG
from langchain_community.document_loaders import PyPDFDirectoryLoader, PyPDFLoader
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_google_vertexai import ChatVertexAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter

# DOCX
try:
    from docx import Document
except Exception:
    Document = None


# =============================================================================
# Flask
# =============================================================================
app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

# =============================================================================
# Config / Env
# =============================================================================
PROJECT_ID = os.environ.get("GCP_PROJECT", "aivle-team0721")
LOCATION   = os.environ.get("GCP_LOCATION", "us-central1")
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
LOCAL_KEY_FILE = os.path.join(BASE_DIR, "firebase-adminsdk.json")

# 업로드/출력 폴더
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Vector DB root (Chroma 1.x 포맷)
RAG_ROOT = os.environ.get("CHROMA_BASE", os.path.join(BASE_DIR, "판례"))
LAW_VECTOR_DB_MAP = {
    "insurance": os.path.join(RAG_ROOT, "보험"),
    "deposit":   os.path.join(RAG_ROOT, "예금"),
    "loan":      os.path.join(RAG_ROOT, "대출"),
}

# 카테고리 정규화
_ALIAS = {
    "insurance": {"insurance","보험","보험(암 포함)","보험(일반)","암보험"},
    "deposit":   {"deposit","예금","적금","savings"},
    "loan":      {"loan","대출"},
}
_ALIAS_FLAT = {v.strip().lower().replace(" ", ""): k
               for k, vs in _ALIAS.items() for v in vs}

def normalize_category(cat: str|None) -> str|None:
    if not cat:
        return None
    return _ALIAS_FLAT.get(cat.strip().lower().replace(" ", ""))

# 검색/게이트 파라미터
TOP_K_DEFAULT      = int(os.environ.get("ANALYZE_TOP_K", "6"))
MAX_QUERY_CHARS    = int(os.environ.get("MAX_QUERY_CHARS", "1500"))
GATE_MIN_SCORE     = int(os.environ.get("GATE_MIN_SCORE", "70"))      # 게이트 통과 점수
MAX_FLAGGED_ABS    = int(os.environ.get("MAX_FLAGGED_ABS", "15"))     # 최대 표출 절대 개수
MAX_FLAGGED_RATIO  = float(os.environ.get("MAX_FLAGGED_RATIO", "0.3"))# 전체 대비 비율

# 조항 필터 문턱값(목차/가짜 조항 제거용)
CLAUSE_MIN_BODY_CHARS   = int(os.environ.get("CLAUSE_MIN_BODY_CHARS", "80"))
CLAUSE_MIN_HANGUL_CHARS = int(os.environ.get("CLAUSE_MIN_HANGUL_CHARS", "30"))

# 설명성 타이틀 힌트
SAFE_TITLE_HINTS = ("목적", "정의", "용어의 정의", "용어의정의", "약관의 명칭", "약관의명칭")

# =============================================================================
# Vertex / LLM / Embeddings
# =============================================================================
credentials = None
try:
    secret_client = secretmanager.SecretManagerServiceClient()
    secret_name   = f"projects/{PROJECT_ID}/secrets/firebase-adminsdk/versions/latest"
    payload = secret_client.access_secret_version(name=secret_name).payload.data.decode("utf-8")
    credentials = service_account.Credentials.from_service_account_info(json.loads(payload))
    logging.info("[BOOT] Secret Manager 자격증명 로드 성공")
except Exception as e:
    logging.warning(f"[BOOT] Secret Manager 실패 → 로컬 키 사용 시도: {e}")
    try:
        if not os.path.exists(LOCAL_KEY_FILE):
            raise FileNotFoundError("로컬 서비스 계정 키 없음: " + LOCAL_KEY_FILE)
        credentials = service_account.Credentials.from_service_account_file(LOCAL_KEY_FILE)
        logging.info("[BOOT] 로컬 자격증명 로드 성공")
    except Exception as file_e:
        logging.error(f"[BOOT] 자격증명 초기화 실패: {file_e}")
        credentials = None

llm = None
embedding_model = None
if credentials:
    try:
        vertexai.init(project=PROJECT_ID, location=LOCATION, credentials=credentials)
        llm = ChatVertexAI(
            model_name="gemini-2.5-flash-lite",
            project=PROJECT_ID,
            location=LOCATION,
            credentials=credentials,
            temperature=0.2,
            max_output_tokens=8192,
        )
        embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        logging.info("[BOOT] Vertex AI 및 임베딩 초기화 성공")
    except Exception:
        logging.exception("[BOOT] Vertex/LLM 초기화 실패")
else:
    logging.error("[BOOT] 자격증명 없음 → LLM/임베딩 사용 불가")

# =============================================================================
# Prompts / Chains
# =============================================================================
# 게이트 프롬프트(중괄호 이스케이프)
gate_prompt = ChatPromptTemplate.from_template(r"""
너는 기업 약관의 '리스크 여부'를 **매우 보수적으로** 판정하는 심사관이다.
기본 원칙:
- 디폴트는 NO(리스크 아님)이다.
- 아래 '중대 신호' 중 **두 가지 이상이 조항의 원문에서 직접 확인될 때만** YES.
- '정의/목적/명칭' 등 설명성 조항은 기본 NO. 다만 정의 자체가 모호해 분쟁 위험이 높은 경우에만 YES.
- 판정은 JSON만 출력한다.

중대 신호(예시):
- 모호/가변 표현: "적절한", "상당한", "필요시", "기타 사유", "합리적인 판단", "등"으로 끝나는 포괄 규정
- 책임/주체 불명확: 누구의 의무인지 불분명, 주체가 바뀌거나 혼재
- 일방적 변경/해지/면책: 사업자에게 일방 재량 부여, 소비자에게 불리
- 핵심 요소 누락: 기한/금액/절차 미기재, 통지 방식 불명확
- 법령/고시/설명의무 위반 소지: 법정고지/철회권/분쟁처리 절차 누락 또는 축소

출력(JSON만):
{{"is_risky": true|false, "score": 0..100, "reasons": ["간결 사유1","간결 사유2"], "evidence": ["원문 발췌1","원문 발췌2"]}}

판정 대상 조항(제목 포함):
{title}

본문:
{clause}
""")
gate_chain = (gate_prompt | llm | StrOutputParser()) if llm else None

# 본 분석 프롬프트
judgment_prompt = ChatPromptTemplate.from_template("""
당신은 기업에서 약관 리스크를 전문적으로 점검하는 법률 전문가입니다.

입력:
- 검토 대상 약관 조항({clause})
- 참고 컨텍스트({similar_cases})  # 내부 판단 참고용

필수 규칙:
1) 약관 조항 중 리스크(모호함, 명확성 부족, 설명의무 위반, 오탈자, 주체 혼동 등)가 있는 부분만 결과를 내세요.
2) 문서에 나온 순서(제1조→제2조→…)로 처리하세요.
3) 한 조항 안에 여러 문제가 있으면 각 문제를 별도 항목으로 나누고, 앞에 1., 2., 3.처럼 번호를 붙여 출력하세요.
4) 아무런 리스크가 없으면 그 조항은 출력하지 마세요.
5) 각 항목의 첫 줄은 반드시 ‘조항 본문에서 발췌한 핵심 문장’만(제목 금지).
6) 결과는 순수 텍스트. 불릿/머리말/요약 금지.
7) 각 항목 내부는 빈 줄 1개로 구분:
   - [문제가 되는 조항]
   - 이유:
   - 수정 제안:

출력 형식(한 항목당 2~3줄), 여러 문제가 있으면 1., 2.로 번호를 붙여 반복:
1. [문제가 되는 조항] 원문 일부/핵심 문장

이유: 무엇이 왜 문제인지(모호·불명확·설명의무 위반 등). 필요한 경우 구체 수치/기한/기준 포함

수정 제안: 소비자가 즉시 이해할 수 있도록 구체 문구로 재작성(수치·기한·정의 포함)

아래 입력을 검토해 위 형식으로만 출력하세요.
{clause}

[참고용 컨텍스트(출력 금지)]
{similar_cases}
""")

judgment_chain = (judgment_prompt | llm | StrOutputParser()) if llm else None

# =============================================================================
# Helpers (조항 분할/정제)
# =============================================================================
CLAUSE_HEADER = r"제\s*\d+\s*조"
CLAUSE_REGEX  = re.compile(rf"({CLAUSE_HEADER}\s*(?:\([^)]+\))?[\s\S]*?)(?={CLAUSE_HEADER}|$)")
TITLE_REGEX   = re.compile(rf"^({CLAUSE_HEADER}[^\n\r]*)")

def _split_title_body(block: str):
    m = TITLE_REGEX.search(block)
    if not m:
        return None, block.strip()
    title = m.group(1).strip()
    body  = block[m.end():].strip()
    return title, body

def _count_hangul(s: str) -> int:
    return len(re.findall(r"[가-힣]", s or ""))

def _is_real_body(body: str) -> bool:
    b = (body or "").strip()
    if len(b) < CLAUSE_MIN_BODY_CHARS:
        return False
    if _count_hangul(b) < CLAUSE_MIN_HANGUL_CHARS:
        return False
    return True

def _collapse_by_title_keep_longest(clauses: List[dict]) -> List[dict]:
    best = {}
    order = {}
    for c in clauses:
        t = c.get("title") or ""
        idx = c.get("index", 10**9)
        order.setdefault(t, idx)
        if not re.search(CLAUSE_HEADER, t):
            key = f"__block__{idx}"
            best[key] = c
            continue
        prev = best.get(t)
        if prev is None or len((c.get("body") or "")) > len((prev.get("body") or "")):
            best[t] = c
    kept = list(best.values())
    kept.sort(key=lambda x: x.get("index", 10**9))
    return kept

def split_into_clauses(text: str) -> List[dict]:
    raw = text or ""
    matches = CLAUSE_REGEX.findall(raw)

    clauses = []
    if matches:
        for idx, m in enumerate(matches, start=1):
            t, b = _split_title_body(m)
            clauses.append({
                "index": idx,
                "title": t or f"제{idx}조",
                "content": m.strip(),
                "body": (b or "").strip(),
            })
    else:
        splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=0)
        chunks = splitter.split_text(raw)
        for i, c in enumerate(chunks):
            t, b = _split_title_body(c)
            clauses.append({
                "index": i+1,
                "title": t or f"블록{i+1}",
                "content": c.strip(),
                "body": (b or c).strip(),
            })

    clauses = _collapse_by_title_keep_longest(clauses)
    clauses = [c for c in clauses if _is_real_body(c.get("body", ""))]

    seen, uniq = set(), []
    for c in clauses:
        key = re.sub(r"\s+", "", (c["title"] + "|" + (c["body"][:160] or "")))
        if key in seen:
            continue
        seen.add(key); uniq.append(c)
    return uniq

# =============================================================================
# 파일 로딩/저장 유틸
# =============================================================================
def load_user_text_from_pdf(pdf_dir: str|None = None, pdf_file: str|None = None) -> str:
    if pdf_file and os.path.isfile(pdf_file):
        docs = PyPDFLoader(pdf_file).load()
    elif pdf_dir and os.path.isdir(pdf_dir):
        docs = PyPDFDirectoryLoader(pdf_dir).load()
    else:
        raise FileNotFoundError("pdf_dir 또는 pdf_file 경로가 올바르지 않습니다.")
    return "".join([d.page_content for d in docs])

def read_txt(path: str) -> str:
    for enc in ("utf-8-sig","utf-8","cp949","euc-kr"):
        try:
            with open(path, "r", encoding=enc, errors="ignore") as f:
                return f.read()
        except Exception:
            continue
    with open(path, "r", errors="ignore") as f:
        return f.read()

def read_docx(path: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다. pip install python-docx")
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)

# ---- 파일 형식 스니핑 ----
def _sniff_ext(file_storage) -> str:
    ext = os.path.splitext((file_storage.filename or "").strip())[1].lower()
    if ext:
        return ext
    head = file_storage.stream.read(4)
    file_storage.stream.seek(0)
    if head.startswith(b"%PDF"): return ".pdf"
    if head.startswith(b"PK\x03\x04"): return ".docx"
    return ""

# =============================================================================
# 벡터스토어 유틸
# =============================================================================
def _ensure_1x_vector_dir(path: str):
    index_dir = os.path.join(path, "index")
    sqlite_file = os.path.join(path, "chroma.sqlite3")
    if os.path.isdir(index_dir) and not os.path.exists(sqlite_file):
        raise RuntimeError(
            "이 벡터DB 폴더는 Chroma 0.4.x(HNSW) 포맷으로 보입니다. "
            "chromadb==1.0.15에서는 사용할 수 없습니다."
        )
    if os.path.isdir(index_dir) and os.path.exists(sqlite_file) and os.environ.get("ALLOW_HNSW_INDEX", "0") != "1":
        try:
            new_name = os.path.join(path, f"index__legacy_{uuid.uuid4().hex[:6]}")
            os.rename(index_dir, new_name)
            logging.warning("[VECTOR] legacy 'index/' 폴더를 %s 로 격리했습니다.", os.path.basename(new_name))
        except Exception as e:
            logging.warning("[VECTOR] 'index/' 격리 실패(계속 진행): %s", e)

def build_vectorstore(path: str) -> Chroma:
    if not os.path.isdir(path):
        raise FileNotFoundError(f"벡터DB 경로가 존재하지 않습니다: {path}")
    _ensure_1x_vector_dir(path)
    logging.info(f"[VECTOR] open (1.x): {os.path.abspath(path)}")
    return Chroma(persist_directory=path, embedding_function=embedding_model)

def _clean_query(q: str) -> str:
    q = re.sub(r"\s+", " ", q or "").strip()
    return q[:MAX_QUERY_CHARS]

def _search_docs(vectorstore: Chroma, query: str, k: int):
    q = _clean_query(query)
    try:
        return list(vectorstore.similarity_search(q, k=k))
    except Exception as e:
        logging.error(f"[VECTOR] similarity_search 실패: {e}")
        return []

# =============================================================================
# 판례 추출 보강 유틸
# =============================================================================
def _strip_html(s: str) -> str:
    s = s or ""
    s = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    s = re.sub(r"<[^>]+>", " ", s)
    return re.sub(r"\s+\n", "\n", s).strip()

def _norm_key(k: str) -> str:
    return re.sub(r"[\s_\.]", "", str(k or "")).lower()

def _lookup_meta(md: dict, candidates) -> str|None:
    if not md:
        return None
    wanted = {_norm_key(x) for x in candidates}
    stack = [md]
    while stack:
        cur = stack.pop()
        if isinstance(cur, dict):
            for k, v in cur.items():
                nk = _norm_key(k)
                if nk in wanted and v not in (None, "", "미상"):
                    return str(v)
                if isinstance(v, (dict, list)):
                    stack.append(v)
        elif isinstance(cur, list):
            stack.extend(cur)
    return None

CASE_NO_RE = re.compile(r"\b(\d{4}\s*[가-힣]{1,3}\s*\d{1,6})\b")
COURT_RE   = re.compile(r"(대법원|고등법원|고법|지방법원|지법|[가-힣]{2,10}(?:법원|고법|지법))")
DATE_RE    = re.compile(r"(\d{4}\.\s*\d{1,2}\.\s*\d{1,2}\.?)")

def _extract_citation_fields(doc):
    md_raw = getattr(doc, "metadata", {}) or {}
    text_raw = (getattr(doc, "page_content", None) or "")
    text = _strip_html(text_raw)

    src = None
    for k in ("source","path","file","filepath","filename","doc_path"):
        if md_raw.get(k):
            src = os.path.basename(str(md_raw[k]))
            break

    court = _lookup_meta(md_raw, ["court","법원","법원명","court_name"])
    date  = _lookup_meta(md_raw, ["decision_date","date","선고일","선고일자","판결선고일","date_str"])
    cno   = _lookup_meta(md_raw, ["case_number","case_no","caseno","사건번호"])
    title = _lookup_meta(md_raw, ["title","사건명","판결명","판례명","name","case_name","document_title","doc_title"])

    if src and not cno:
        m = CASE_NO_RE.search(src)
        if m: cno = m.group(1)
    if src and not title:
        base = os.path.splitext(src)[0]
        cand = re.sub(r"[_\-\[\]\(\)]", " ", base).strip()
        if len(cand) >= 2:
            title = cand

    if not court:
        m = COURT_RE.search(text); court = m.group(1) if m else None
    if not date:
        m = DATE_RE.search(text);  date  = m.group(1) if m else None
    if not cno:
        m = CASE_NO_RE.search(text); cno = m.group(1) if m else None

    if not title:
        for line in text.splitlines():
            t = line.strip().strip("【】[]")
            if 2 <= len(t) <= 60 and any(k in t for k in ("사건","판결","취소","보험","대출")):
                title = t; break

    return {
        "court": court or "미상",
        "date":  date  or "미상",
        "case_no": cno or "미상",
        "title": title or "미상",
    }

def _doc_to_citation_snippet(doc) -> str:
    info = _extract_citation_fields(doc)
    header_parts = []
    if info["court"]  != "미상": header_parts.append(f"법원: {info['court']}")
    if info["date"]   != "미상": header_parts.append(f"선고일: {info['date']}")
    if info["case_no"]!= "미상": header_parts.append(f"사건번호: {info['case_no']}")
    if info["title"]  != "미상": header_parts.append(f"사건명: {info['title']}")
    header = " | ".join(header_parts)
    page = _strip_html(getattr(doc, "page_content", "") or "")
    return (header + "\n" if header else "") + page

def _citations_catalog_from_docs(docs: List) -> str:
    seen, rows = set(), []
    for d in docs:
        info = _extract_citation_fields(d)
        key = (info["court"], info["date"], info["case_no"], info["title"])
        if key in seen:
            continue
        seen.add(key)
        rows.append(f"- 법원: {info['court']} | 선고일: {info['date']} | 사건번호: {info['case_no']} | 사건명: {info['title']}")
    return "\n".join(rows)

# =============================================================================
# 게이트(사전 판정)
# =============================================================================
def _extract_json_block(s: str) -> dict:
    try:
        m = re.search(r"\{[\s\S]*\}", s or "")
        if m:
            s = m.group(0)
        return json.loads(s)
    except Exception:
        return {}

def is_descriptive_title(title: str|None) -> bool:
    t = (title or "").replace(" ", "")
    return any(h in t for h in SAFE_TITLE_HINTS)

def gate_clause(title: str, body: str) -> dict:
    if not llm or not gate_chain or not _is_real_body(body):
        return {"is_risky": False, "score": 0, "reasons": [], "evidence": []}
    hint = is_descriptive_title(title)
    try:
        out = gate_chain.invoke({"title": title, "clause": body}) or ""
        data = _extract_json_block(out)
        is_risky = bool(data.get("is_risky"))
        score = int(data.get("score") or 0)
        reasons = data.get("reasons") or []
        evidence = data.get("evidence") or []
        if hint and score > 0:
            score = max(0, score - 20)  # 설명성 타이틀 감점
        if is_risky and (len(evidence) < 1 or len(reasons) < 1):
            is_risky, score = False, 0
        return {"is_risky": is_risky, "score": score, "reasons": reasons, "evidence": evidence}
    except Exception:
        logging.exception("[GATE] 게이트 판정 실패")
        return {"is_risky": False, "score": 0, "reasons": [], "evidence": []}

# =============================================================================
# 본 분석(수정 제안)
# =============================================================================
def analyze_single_clause(clause_text: str, vectorstore: Chroma,
                          top_k: int = TOP_K_DEFAULT) -> str:
    if not _is_real_body(clause_text):
        return ""
    docs = _search_docs(vectorstore, clause_text, k=top_k)
    if not docs:
        return ""

    similar_text = "\n\n".join([_doc_to_citation_snippet(d) for d in docs])

    out = judgment_chain.invoke({
        "clause": clause_text,
        "similar_cases": similar_text,
    }) if judgment_chain else ""
    out = (out or "").strip()

    # 안전 가드(선택): 혹시 모델이 '관련 판례:'를 찍어도 제거
    if out:
        cleaned = []
        for ln in out.splitlines():
            if ln.strip().startswith("관련 판례:"):
                continue
            cleaned.append(ln)
        out = "\n".join(cleaned).strip()
    return out

# 분석결과 → (원문 스니펫, 대체 문구) 페어 추출
PAIR_BLOCK_RE = re.compile(
    r"^\s*(?:\d+\s*[\.\)]\s*)?\[문제가 되는 조항\]\s*(?P<src>.+?)\s*[\r\n]+"
    r"(?:설명|문제\s*이유)\s*:\s*(?P<reason>.+?)\s*[\r\n]+"
    r"수정\s*제안\s*:\s*(?P<dst>.+?)$",
    re.DOTALL | re.MULTILINE
)

def parse_replacement_pairs(analysis_text: str) -> List[Tuple[str, str]]:
    pairs: List[Tuple[str, str]] = []
    text = analysis_text or ""
    for m in PAIR_BLOCK_RE.finditer(text):
        src = (m.group("src") or "").strip()
        dst = (m.group("dst") or "").strip()
        if src and dst:
            pairs.append((src, dst))
    return pairs

def _normalize_space(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

def replace_text_loose(haystack: str, needle: str, repl: str) -> Tuple[str, int]:
    if not haystack or not needle:
        return haystack, 0
    if needle in haystack:
        return haystack.replace(needle, repl), haystack.count(needle)
    base = _normalize_space(needle)
    if not base:
        return haystack, 0
    pat = re.compile(re.escape(base).replace(r"\ ", r"\s+"))
    cnt = 0
    def _sub(_m):
        nonlocal cnt
        cnt += 1
        return repl
    new_text, n = pat.subn(_sub, _normalize_space(haystack))
    if n:
        return new_text, cnt
    return haystack, 0

def apply_pairs_to_docx(docx_path: str, pairs: List[Tuple[str, str]]) -> Tuple[str, int]:
    if Document is None:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다.")
    doc = Document(docx_path)
    applied = 0

    def _apply_to_paragraph(p):
        nonlocal applied
        text = p.text
        for src, dst in pairs:
            new_text, n = replace_text_loose(text, src, dst)
            if n:
                text = new_text
                applied += n
        if text != p.text:
            p.text = text

    for p in doc.paragraphs:
        _apply_to_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_to_paragraph(p)

    out_name = f"applied_{uuid.uuid4().hex[:8]}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_name, applied

def build_docx_from_text(text: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx가 설치되어 있지 않습니다.")
    doc = Document()
    for line in (text or "").splitlines():
        doc.add_paragraph(line)
    out_name = f"text_{uuid.uuid4().hex[:8]}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_name

def apply_pairs_to_text(text: str, pairs: List[Tuple[str, str]]) -> Tuple[str, int]:
    applied = 0
    out = text or ""
    for src, dst in pairs:
        out, n = replace_text_loose(out, src, dst)
        applied += n
    return out, applied

# =============================================================================
# 공통 실행 루틴: 게이트 → 상위만 본분석
# =============================================================================
def _analyze_clauses_with_gate(clauses: List[dict], vectorstore: Chroma):
    gated = []
    for c in clauses:
        body = (c.get("body") or "").strip()
        title = (c.get("title") or "").strip()
        if not _is_real_body(body):
            continue
        g = gate_clause(title, body)
        if g.get("is_risky") and g.get("score", 0) >= GATE_MIN_SCORE:
            gated.append({**c, "gate": g})

    if not gated:
        return [], "", []

    cap_by_ratio = max(1, int(len(clauses) * MAX_FLAGGED_RATIO))
    cap = min(MAX_FLAGGED_ABS, cap_by_ratio)
    gated.sort(key=lambda x: x["gate"]["score"], reverse=True)
    selected = gated[:cap]

    results = []
    for c in selected:
        analysis = analyze_single_clause(c["body"], vectorstore, top_k=TOP_K_DEFAULT)
        if analysis:
            results.append({"index": c["index"], "title": c["title"], "analysis": analysis, "gate": c["gate"]})

    joined = "\n\n".join([r["analysis"] for r in results])
    pairs = parse_replacement_pairs(joined)
    return results, joined, pairs

# =============================================================================
# API
# =============================================================================
@app.route("/api/health", methods=["GET"])
def health():
    return {"ok": True, "service": "analyze_terms", "time": datetime.now(timezone.utc).isoformat()}

@app.route("/api/download/<path:filename>", methods=["GET"])
def download_file(filename):
    safe = os.path.basename(filename)
    return send_from_directory(OUTPUT_DIR, safe, as_attachment=True)

@app.route("/__whoami", methods=["GET"])
def whoami():
    routes = sorted([str(r) for r in app.url_map.iter_rules()])
    return {
        "file": __file__,
        "cwd": os.getcwd(),
        "python": sys.executable,
        "routes": routes[:200],
    }

@app.route("/api/debug/vector_db", methods=["GET"])
def debug_vector_db():
    info = {}
    for k, p in LAW_VECTOR_DB_MAP.items():
        p_abs = os.path.abspath(p)
        version = None
        vfile = os.path.join(p_abs, "VERSION")
        if os.path.exists(vfile):
            try:
                with open(vfile, "r", encoding="utf-8") as f:
                    version = f.read().strip()
            except Exception:
                pass
        info[k] = {
            "path": p_abs,
            "exists": os.path.isdir(p_abs),
            "has_sqlite": os.path.exists(os.path.join(p_abs, "chroma.sqlite3")),
            "has_index_dir": os.path.exists(os.path.join(p_abs, "index")),
            "version_file": version,
            "sample": sorted(os.listdir(p_abs))[:12] if os.path.isdir(p_abs) else [],
        }
    return info

@app.after_request
def _after(resp):
    resp.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization,x-authenticated-user-uid')
    resp.headers.add('Access-Control-Allow-Methods', 'GET,POST,PUT,DELETE,OPTIONS')
    return resp

# JSON 본문 분석(파일 없이)
@app.route("/api/analyze-terms", methods=["POST", "OPTIONS"])
@cross_origin(origin="*")
def analyze_terms():
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200

    data = request.get_json(silent=True) or {}
    raw_text = data.get("text", "")
    category_raw = data.get("category", "")
    category = normalize_category(category_raw)
    vector_db_path = data.get("vector_db_path") or (LAW_VECTOR_DB_MAP.get(category) if category else None)
    limit = int(data.get("limit", 0))

    try:
        full_text = (raw_text or "").strip()
        if not full_text:
            pdf_dir = data.get("pdf_dir")
            pdf_file = data.get("pdf_file")
            full_text = load_user_text_from_pdf(pdf_dir=pdf_dir, pdf_file=pdf_file)

        if not vector_db_path:
            return jsonify({"ok": False, "error": f"category가 유효하지 않습니다: {category_raw}. 허용: {list(LAW_VECTOR_DB_MAP)}"}), 400

        clauses = split_into_clauses(full_text)
        if limit > 0:
            clauses = clauses[:limit]

        vectorstore = build_vectorstore(vector_db_path)

        results, joined, pairs = _analyze_clauses_with_gate(clauses, vectorstore)

        return jsonify({
            "ok": True,
            "category": category,
            "vector_db_path": vector_db_path,
            "count_clauses": len(clauses),
            "count_flagged": len(results),
            "results": results,
            "text": joined,
            "pairs": [{"from": s, "to": d} for s, d in pairs],
        })
    except Exception as e:
        logging.exception("[API] /api/analyze-terms 오류")
        return jsonify({"ok": False, "error": str(e)}), 500

# 업로드 파일 분석 + 원문에 수정 적용하여 새 파일 제공
@app.route("/api/analyze-terms-upload", methods=["POST", "OPTIONS"])
@cross_origin(origin="*")
def analyze_terms_upload():
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "file 필드가 없습니다."}), 400

    f = request.files["file"]
    if not f or (f.filename or "").strip() == "":
        return jsonify({"ok": False, "error": "파일이 비어 있습니다."}), 400

    category_raw = request.form.get("category", "")
    category = normalize_category(category_raw)
    limit = int(request.form.get("limit", 0))

    ext = _sniff_ext(f)
    safe_name = secure_filename(f.filename or f"upload{ext or ''}")
    src_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex}{ext or ''}")

    try:
        f.save(src_path)
        logging.info(f"[UPLOAD] {safe_name} → {src_path}")

        if ext == ".pdf":
            full_text = load_user_text_from_pdf(pdf_file=src_path)
        elif ext == ".docx":
            full_text = read_docx(src_path)
        elif ext == ".txt" or ext == "":
            full_text = read_txt(src_path)
        else:
            return jsonify({"ok": False, "error": f"지원하지 않는 형식입니다: {ext or '(알 수 없음)'} . txt/pdf/docx만 허용"}), 400

        if not category or category not in LAW_VECTOR_DB_MAP:
            return jsonify({"ok": False, "error": f"category가 유효하지 않습니다: {category_raw}. 허용: {list(LAW_VECTOR_DB_MAP)}"}), 400

        vector_db_path = LAW_VECTOR_DB_MAP[category]

        clauses = split_into_clauses(full_text)
        if limit > 0:
            clauses = clauses[:limit]

        vectorstore = build_vectorstore(vector_db_path)

        results, joined, pairs = _analyze_clauses_with_gate(clauses, vectorstore)

        output_filename, applied_count = None, 0
        try:
            if ext == ".docx":
                output_filename, applied_count = apply_pairs_to_docx(src_path, pairs)
            else:
                replaced_text, applied_count = apply_pairs_to_text(full_text, pairs)
                output_filename = build_docx_from_text(replaced_text)
        except Exception:
            logging.exception("치환 적용/저장 실패")
            output_filename, applied_count = None, 0

        download_url = None
        if output_filename:
            download_url = request.host_url.rstrip('/') + f"/api/download/{output_filename}"

        return jsonify({
            "ok": True,
            "category": category,
            "vector_db_path": vector_db_path,
            "count_clauses": len(clauses),
            "count_flagged": len(results),
            "results": results,
            "text": joined,
            "pairs": [{"from": s, "to": d} for s, d in pairs],
            "applied_replacements": applied_count,
            "output_file": output_filename,
            "output_url": download_url,
        })
    except Exception as e:
        logging.exception("[API] /api/analyze-terms-upload 오류")
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        if os.path.exists(src_path):
            try:
                os.remove(src_path)
                logging.info(f"[CLEANUP] 임시 파일 삭제: {src_path}")
            except Exception as e_clean:
                logging.warning(f"[CLEANUP] 임시 파일 삭제 실패: {src_path}, error: {e_clean}")

# =============================================================================
# Run (리로더 끔: 중복 프로세스/포트 혼선 방지)
# =============================================================================
if __name__ == "__main__":
    port = int(os.environ.get("PORT", os.environ.get("PY_PORT", 8080)))
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)
